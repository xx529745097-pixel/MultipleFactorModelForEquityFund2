import copy
import html
from io import BytesIO, IOBase
from typing import BinaryIO, List, Literal, Union

import altair as alt
import dataframe_image as dfi
import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image
from pandas.io.formats.style import Styler
from docx import Document
from docx.shared import Cm, Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

##########
# 添加标题
##########
def doc_addHead(
    doc,  # Document 对象
    level,  # int, 标题级数
    text,  # 标题内容
    font=u"黑体",  # 字体类型
    size=16  # 字体大小
):
    title = doc.add_heading(level=level)
    title_run = title.add_run(text)
    title_run.font.size = Pt(size)
    title_run.font.name = font
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), font)  # 设置中文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.bold = True
    return doc

###########
# 添加正文
###########
def doc_addPara(
    doc,  # Document 对象
    text,  # 段落文本内容
    font=u"仿宋_GB2312",  # 字体类型
    size=12,  # 字体大小
    bold=False,  # 字体加粗
    first_line_indent=True  # 首行缩进
):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0  # 行间距倍数
    p.paragraph_format.first_line_indent = Pt(size) * 2 if first_line_indent else 0
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # 设置中文字体
    r.font.bold = bold
    return doc

#######################################
# 添加图片，可通过width_scale设置图片缩放比例
#######################################
def doc_addPic(
    doc: Document,  # Document 对象
    fig: Union[
        bytes,
        str,
        plt.Figure,
        BinaryIO,
        alt.Chart,
        pd.DataFrame,
        Styler,
        List[List[Union[float, str]]],
    ],  # 支持图片类型
    width_scale: float = 1.0,  # 图片相对文档宽度缩放倍数(不含页边距)，默认为1倍文档内容宽度
):
    page_content_width_inches = doc.sections[0].page_width.inches - doc.sections[0].left_margin.inches - doc.sections[0].right_margin.inches
    # prepare figure
    if isinstance(fig, str) or isinstance(fig, IOBase):
        figio = fig
    elif isinstance(fig, bytes):
        figio = BytesIO(fig)
    elif isinstance(fig, plt.Figure):
        figio = BytesIO()
        fig.set_size_inches(12, 8)
        fig.savefig(figio, format="png", bbox_inches="tight", dpi=200)
    elif isinstance(fig, alt.VegaLiteSchema):
        figio = BytesIO()
        fig.properties(width=480, height=270).save(figio, format="png", scale_factor=2)
    elif (
        isinstance(fig, List)
        or isinstance(fig, pd.DataFrame)
        or isinstance(fig, Styler)
    ):
        if isinstance(fig, List):
            fig = pd.DataFrame(fig)
        figio = BytesIO()
        dfi.export(fig, figio, dpi=160)
    else:
        raise ValueError(f"{type(fig)} {repr(fig)} is not supported")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("")
    # add pic and adjust image size
    r.add_picture(figio, width=width_scale * Inches(page_content_width_inches))
    return doc

#####################
# 按顺序编号各级标题
#####################
def doc_processHead(doc):
    head1 = 0
    head2 = 0
    head3 = 0
    head4 = 0

    for para in doc.paragraphs:
        if para.text == "报告摘要":
            continue
        style_name = para.style.name
        if style_name == "Heading 1":
            head1 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text, str(head1) + " " + para.text
                )
            head2 = 0
            head3 = 0
            head4 = 0
        if style_name == "Heading 2":
            head2 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text, str(head1) + "." + str(head2) + " " + para.text
                )
            head3 = 0
            head4 = 0
        if style_name == "Heading 3":
            head3 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text,
                    str(head1) + "." + str(head2) + "." + str(head3) + " " + para.text
                )
            head4 = 0
        if style_name == "Heading 4":
            head4 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text,
                    str(head1) + "." + str(head2) + "." + str(head3) + "." + str(head4) + " " + para.text
                )
    return doc

##########################
# header内容替换(保留源格式)
##########################
def doc_replaceHeaderText(
    doc,  # Document obj
    text_like,  # str, 被替换的字符串
    str_replace_with  # str, 填入的字符串
):
    # header.is_linked_to_previous默认为True, 因此仅需修改首个header
    header = doc.sections[0].header
    for para in header.paragraphs:
        for run in para.runs:
            if text_like in run.text:
                run.text = run.text.replace(text_like, str_replace_with)
    return doc

#######################
# body内容替换(保留源格式)
#######################
def doc_replaceBodyText(
    doc,  # Document obj
    text_like,  # str, 被替换的字符串
    str_replace_with  # str, 填入的字符串
):
    for para in doc.paragraphs:
        for run in para.runs:
            if text_like in run.text:
                run.text = run.text.replace(text_like, str_replace_with)
    return doc
